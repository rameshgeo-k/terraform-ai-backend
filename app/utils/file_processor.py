"""
File Processing Utilities
Extract text content from various file formats
"""

import logging
import mimetypes
from pathlib import Path
from typing import Optional, Dict, Any
import io

logger = logging.getLogger(__name__)


class FileProcessor:
    """Process and extract text from various file formats"""
    
    # Supported file extensions and their MIME types
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.py': 'text/x-python',
        '.js': 'text/javascript',
        '.ts': 'text/typescript',
        '.java': 'text/x-java',
        '.cpp': 'text/x-c++',
        '.c': 'text/x-c',
        '.go': 'text/x-go',
        '.rs': 'text/x-rust',
        '.tf': 'text/x-terraform',
        '.json': 'application/json',
        '.xml': 'application/xml',
        '.yaml': 'text/yaml',
        '.yml': 'text/yaml',
        '.sh': 'application/x-sh',
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.doc': 'application/msword',
        '.csv': 'text/csv',
        '.html': 'text/html',
        '.htm': 'text/html',
    }
    
    @staticmethod
    def is_supported(filename: str) -> bool:
        """Check if file type is supported"""
        ext = Path(filename).suffix.lower()
        return ext in FileProcessor.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def get_mime_type(filename: str) -> Optional[str]:
        """Get MIME type for file"""
        ext = Path(filename).suffix.lower()
        return FileProcessor.SUPPORTED_EXTENSIONS.get(ext)
    
    @staticmethod
    async def extract_text(
        file_content: bytes,
        filename: str,
        mime_type: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Extract text content from file
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            mime_type: Optional MIME type override
            
        Returns:
            Dictionary with 'text', 'metadata', and 'format' keys
        """
        if mime_type is None:
            mime_type = FileProcessor.get_mime_type(filename)
        
        ext = Path(filename).suffix.lower()
        
        try:
            # Text-based files
            if mime_type and mime_type.startswith('text/') or ext in ['.txt', '.md', '.py', '.js', '.ts', '.java', '.cpp', '.c', '.go', '.rs', '.tf', '.json', '.xml', '.yaml', '.yml', '.sh', '.html', '.htm', '.csv']:
                text = file_content.decode('utf-8', errors='ignore')
                return {
                    'text': text,
                    'format': ext[1:] if ext else 'text',
                    'metadata': {
                        'filename': filename,
                        'size': len(file_content),
                        'type': 'text'
                    }
                }
            
            # PDF files
            elif ext == '.pdf':
                text = await FileProcessor._extract_pdf(file_content)
                return {
                    'text': text,
                    'format': 'pdf',
                    'metadata': {
                        'filename': filename,
                        'size': len(file_content),
                        'type': 'document'
                    }
                }
            
            # DOCX files
            elif ext == '.docx':
                text = await FileProcessor._extract_docx(file_content)
                return {
                    'text': text,
                    'format': 'docx',
                    'metadata': {
                        'filename': filename,
                        'size': len(file_content),
                        'type': 'document'
                    }
                }
            
            # DOC files (older Word format)
            elif ext == '.doc':
                # Basic text extraction for .doc files (limited support)
                text = file_content.decode('utf-8', errors='ignore')
                return {
                    'text': text,
                    'format': 'doc',
                    'metadata': {
                        'filename': filename,
                        'size': len(file_content),
                        'type': 'document',
                        'warning': 'Limited support for .doc format'
                    }
                }
            
            else:
                raise ValueError(f"Unsupported file format: {ext}")
        
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            raise ValueError(f"Failed to extract text from {filename}: {str(e)}")
    
    @staticmethod
    async def _extract_pdf(file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
            
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_parts.append(page.extract_text())
            
            return '\n\n'.join(text_parts)
        
        except ImportError:
            raise ValueError("PyPDF2 not installed. Run: pip install PyPDF2")
        except Exception as e:
            raise ValueError(f"Failed to extract PDF text: {str(e)}")
    
    @staticmethod
    async def _extract_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            import docx
            
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return '\n\n'.join(text_parts)
        
        except ImportError:
            raise ValueError("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            raise ValueError(f"Failed to extract DOCX text: {str(e)}")
    
    @staticmethod
    def validate_file_size(file_size: int, max_size_mb: int = 10) -> bool:
        """
        Validate file size
        
        Args:
            file_size: File size in bytes
            max_size_mb: Maximum allowed size in MB
            
        Returns:
            True if valid, False otherwise
        """
        max_size_bytes = max_size_mb * 1024 * 1024
        return file_size <= max_size_bytes
    
    @staticmethod
    def get_file_info(filename: str, file_size: int) -> Dict[str, Any]:
        """Get file information"""
        ext = Path(filename).suffix.lower()
        mime_type = FileProcessor.get_mime_type(filename)
        
        return {
            'filename': filename,
            'extension': ext,
            'mime_type': mime_type,
            'size': file_size,
            'size_mb': round(file_size / (1024 * 1024), 2),
            'supported': FileProcessor.is_supported(filename)
        }
